"""Kiểm thử tầng bóc tách phiếu khám.

Toàn bộ test chạy được offline, không cần OPENAI_API_KEY: phần gọi model được
thay bằng bộ giả lập, còn phần quan trọng nhất — kiểm & ép kiểu kết quả — vốn là
code thuần nên test trực tiếp.
"""

from __future__ import annotations

import base64
from pathlib import Path

import pytest

from app.modules.documents.extractor.file_reader import (
    UnsupportedFileError,
    read_document,
)
from app.modules.documents.extractor.schema_spec import (
    KIND_MANY_OF,
    KIND_NUMBER,
    KIND_ONE_OF,
    build_field_specs,
    coerce_extraction,
    count_filled,
    render_spec_for_prompt,
)
from app.modules.documents.extractor.stub import StubExamExtractor
from app.modules.patients.form_service import FormService


@pytest.fixture(scope="module")
def f20_block():
    schema = FormService().get_schema("f20")
    return next(b for b in schema["blocks"] if b["id"] == "examination")


@pytest.fixture(scope="module")
def f32_block():
    schema = FormService().get_schema("f32")
    return next(b for b in schema["blocks"] if b["id"] == "examination")


@pytest.fixture(scope="module")
def f20_specs(f20_block):
    return build_field_specs(f20_block)


# ==================== Sinh đặc tả từ schema ====================


def test_specs_cover_every_leaf_and_matrix_row(f20_specs):
    paths = {spec.path for spec in f20_specs}
    assert "general.weight" in paths
    assert "cardiovascular.finding" in paths
    assert "general_lab.blood_count.rbc" in paths
    assert "treatment.inpatient_days" in paths
    # Mỗi dòng của bảng Ý thức phải thành một trường riêng.
    assert "mental_exam.consciousness.orientation.space" in paths
    assert "mental_exam.consciousness.orientation.surroundings" in paths


def test_specs_differ_between_diseases(f20_block, f32_block):
    f20 = {s.path for s in build_field_specs(f20_block)}
    f32 = {s.path for s in build_field_specs(f32_block)}
    # Bảng triệu chứng cơ thể chỉ có ở mẫu trầm cảm.
    assert any(p.startswith("mental_exam.somatic.") for p in f32)
    assert not any(p.startswith("mental_exam.somatic.") for p in f20)
    # Ngược lại, khám tâm thần F20 có mục hoang tưởng riêng.
    assert "mental_exam.thought.content" in f20
    assert "mental_exam.thought.content" not in f32


def test_spec_kinds_are_mapped(f20_specs):
    by_path = {spec.path: spec for spec in f20_specs}
    assert by_path["general.weight"].kind == KIND_NUMBER
    assert by_path["general.weight"].unit == "kg"
    assert by_path["mental_exam.general_observation.contact"].kind == KIND_ONE_OF
    assert by_path["mental_exam.general_observation.attitude"].kind == KIND_MANY_OF
    # Ô trong bảng Có/Không là chọn-1 với đúng 2 lựa chọn.
    orientation = by_path["mental_exam.consciousness.orientation.space"]
    assert orientation.kind == KIND_ONE_OF
    assert orientation.option_values == ["normal", "disordered"]


def test_prompt_spec_contains_codes_and_labels(f20_specs):
    text = render_spec_for_prompt(f20_specs)
    assert "mental_exam.general_observation.contact" in text
    assert "limited=Hợp tác hạn chế" in text
    assert "### " in text  # có chia mục


# ==================== Kiểm & ép kiểu kết quả model ====================


def test_coerce_accepts_codes(f20_specs):
    data, warnings = coerce_extraction(
        {
            "general.weight": 58,
            "mental_exam.general_observation.contact": "limited",
            "mental_exam.general_observation.attitude": ["natural", "tense"],
        },
        f20_specs,
    )
    assert data["general"]["weight"] == 58
    assert data["mental_exam"]["general_observation"]["contact"] == "limited"
    assert data["mental_exam"]["general_observation"]["attitude"] == ["natural", "tense"]
    assert warnings == []


def test_coerce_accepts_vietnamese_labels(f20_specs):
    """Model hay trả nhãn thay vì mã — phải nhận cả hai, kể cả khi mất dấu."""
    data, _ = coerce_extraction(
        {
            "mental_exam.general_observation.contact": "Hợp tác hạn chế",
            "mental_exam.intelligence.level": "sa sut tri tue",
            "mental_exam.consciousness.orientation.time": "Rối loạn",
        },
        f20_specs,
    )
    assert data["mental_exam"]["general_observation"]["contact"] == "limited"
    assert data["mental_exam"]["intelligence"]["level"] == "dementia"
    assert data["mental_exam"]["consciousness"]["orientation"]["time"] == "disordered"


def test_coerce_accepts_nested_input(f20_specs):
    """Model có thể trả lồng thay vì phẳng — kết quả phải như nhau."""
    nested, _ = coerce_extraction(
        {"general": {"weight": 58, "blood_pressure": "120/80"}}, f20_specs
    )
    flat, _ = coerce_extraction(
        {"general.weight": 58, "general.blood_pressure": "120/80"}, f20_specs
    )
    assert nested == flat


def test_coerce_strips_units_and_parses_decimals(f20_specs):
    data, _ = coerce_extraction(
        {
            "general.weight": "58 kg",
            "general.temperature": "36,8",
            "general_lab.blood_count.rbc": "4.82 T/L",
        },
        f20_specs,
    )
    assert data["general"]["weight"] == 58
    assert data["general"]["temperature"] == 36.8
    assert data["general_lab"]["blood_count"]["rbc"] == 4.82


def test_coerce_parses_date_formats(f20_specs):
    for raw in ("2026-03-15", "15/03/2026", "15-03-2026"):
        data, _ = coerce_extraction({"exam_info.exam_date": raw}, f20_specs)
        assert data["exam_info"]["exam_date"] == "2026-03-15"


def test_coerce_drops_invalid_option_with_warning(f20_specs):
    """Giá trị không thuộc template thì bỏ hẳn — thà trống còn hơn ghi bậy."""
    data, warnings = coerce_extraction(
        {"mental_exam.general_observation.contact": "bệnh nhân ngủ li bì"}, f20_specs
    )
    assert "general_observation" not in data.get("mental_exam", {})
    assert any("không thuộc danh sách" in w for w in warnings)


def test_negation_never_matches_affirmative_option(f20_specs):
    """Chốt chặn quan trọng nhất: “không bình thường” tuyệt đối không được
    hiểu thành “Bình thường”. Khớp theo chuỗi con dễ đảo ngược nghĩa bệnh án."""
    data, warnings = coerce_extraction(
        {"mental_exam.consciousness.orientation.space": "không bình thường"}, f20_specs
    )
    assert data == {}
    assert warnings

    # Ngược lại, rút gọn nhãn cùng sắc thái thì vẫn nhận.
    data, _ = coerce_extraction(
        {"mental_exam.general_observation.hygiene": "Gọn gàng"}, f20_specs
    )
    assert data["mental_exam"]["general_observation"]["hygiene"] == ["tidy"]


def test_coerce_drops_unknown_path_with_warning(f20_specs):
    data, warnings = coerce_extraction(
        {"khong_ton_tai.abc": "x", "general.weight": 60}, f20_specs
    )
    assert data == {"general": {"weight": 60}}
    assert any("không có trong biểu mẫu" in w for w in warnings)


def test_coerce_ignores_blank_values(f20_specs):
    data, warnings = coerce_extraction(
        {"general.body_condition": "   ", "general.skin_mucosa": None, "general.weight": 60},
        f20_specs,
    )
    assert data == {"general": {"weight": 60}}
    assert warnings == []


def test_coerce_multi_select_from_string(f20_specs):
    data, _ = coerce_extraction(
        {"mental_exam.general_observation.attitude": "Tự nhiên, Căng thẳng"}, f20_specs
    )
    assert data["mental_exam"]["general_observation"]["attitude"] == ["natural", "tense"]


def test_coerce_multi_select_partially_invalid(f20_specs):
    data, warnings = coerce_extraction(
        {"mental_exam.general_observation.attitude": ["natural", "xyz"]}, f20_specs
    )
    assert data["mental_exam"]["general_observation"]["attitude"] == ["natural"]
    assert len(warnings) == 1


def test_coerce_yes_no_table(f32_block):
    specs = build_field_specs(f32_block)
    data, warnings = coerce_extraction(
        {
            "mental_exam.emotion.symptoms.depressed_mood": "Có",
            "mental_exam.emotion.symptoms.anhedonia": "no",
            "mental_exam.emotion.symptoms.suicidal_ideation": "Không",
        },
        specs,
    )
    symptoms = data["mental_exam"]["emotion"]["symptoms"]
    assert symptoms == {"depressed_mood": "yes", "anhedonia": "no", "suicidal_ideation": "no"}
    assert warnings == []


def test_count_filled(f20_specs):
    data, _ = coerce_extraction(
        {"general.weight": 60, "general.height": 170, "diagnosis.definitive": "F20.0"},
        f20_specs,
    )
    assert count_filled(data) == 3


# ==================== Bộ giả lập ====================


def test_stub_extractor_output_survives_coercion(f20_specs):
    """Dữ liệu giả lập phải hợp lệ 100% — nếu không thì bộ ép kiểu đang sai."""
    from app.modules.documents.extractor.file_reader import DocumentContent

    outcome = StubExamExtractor().extract(
        DocumentContent(kind="text", page_count=1, text="x"), f20_specs, "Tâm thần phân liệt"
    )
    data, warnings = coerce_extraction(outcome.raw_result, f20_specs)
    assert warnings == []
    assert count_filled(data) > 20


# ==================== Đọc file ====================


def _make_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("PHẦN 3: KHÁM BỆNH")
    document.add_paragraph("3.1. Toàn thân: Thể trạng trung bình")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Triệu chứng"
    table.cell(0, 1).text = "Có"
    table.cell(0, 2).text = "Không"
    table.cell(1, 0).text = "Khí sắc giảm"
    table.cell(1, 1).text = "x"
    document.save(str(path))


def _make_text_pdf(path: Path) -> None:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 100), "PHAN 3: KHAM BENH - The trang trung binh")
    page.insert_text((72, 130), "Mach 78 l/p, Nhiet do 36.8, Huyet ap 120/80")
    # Cần đủ dài để vượt ngưỡng "có lớp văn bản".
    for index in range(12):
        page.insert_text((72, 160 + index * 16), f"Dong noi dung so {index} de kiem tra lop text")
    document.save(str(path))
    document.close()


def _make_scanned_pdf(path: Path) -> None:
    """PDF chỉ chứa ảnh — không có lớp text, giống bản scan."""
    import fitz

    source = fitz.open()
    page = source.new_page()
    page.draw_rect(fitz.Rect(50, 50, 300, 200), color=(0, 0, 0))
    pixmap = page.get_pixmap(dpi=72)
    source.close()

    document = fitz.open()
    target = document.new_page()
    target.insert_image(target.rect, stream=pixmap.tobytes("png"))
    document.save(str(path))
    document.close()


def test_read_docx_includes_table_cells(tmp_path):
    path = tmp_path / "phieu.docx"
    _make_docx(path)
    content = read_document(path, prefer="text")
    assert content.kind == "text"
    assert "Thể trạng trung bình" in content.text
    assert "Khí sắc giảm" in content.text  # nội dung trong bảng cũng phải lấy được


def test_read_pdf_with_text_layer(tmp_path):
    path = tmp_path / "digital.pdf"
    _make_text_pdf(path)
    content = read_document(path, prefer="text")
    assert content.kind == "text"
    assert "KHAM BENH" in content.text


def test_scanned_pdf_falls_back_to_images(tmp_path):
    """Người dùng bấm 'Upload phiếu khám' nhưng file lại là bản scan."""
    path = tmp_path / "scan.pdf"
    _make_scanned_pdf(path)
    content = read_document(path, prefer="text", dpi=72)
    assert content.kind == "images"
    assert content.images
    assert "bản scan" in content.note


def test_text_pdf_used_directly_even_in_ocr_mode(tmp_path):
    """Bấm OCR nhưng file có sẵn text -> đọc text, vừa chính xác vừa rẻ hơn."""
    path = tmp_path / "digital.pdf"
    _make_text_pdf(path)
    content = read_document(path, prefer="images")
    assert content.kind == "text"
    assert "lớp văn bản" in content.note


def test_image_is_passed_through_as_base64(tmp_path):
    import fitz

    path = tmp_path / "anh.png"
    document = fitz.open()
    page = document.new_page()
    page.get_pixmap(dpi=36).save(str(path))
    document.close()

    content = read_document(path)
    assert content.kind == "images"
    assert base64.b64decode(content.images[0])[:4] == b"\x89PNG"


def test_unsupported_suffix_rejected(tmp_path):
    path = tmp_path / "file.txt"
    path.write_text("noi dung")
    with pytest.raises(UnsupportedFileError):
        read_document(path)


def test_legacy_doc_rejected_with_hint(tmp_path):
    path = tmp_path / "file.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0")
    with pytest.raises(UnsupportedFileError, match="docx"):
        read_document(path)
