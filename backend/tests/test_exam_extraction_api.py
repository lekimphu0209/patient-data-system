"""Kiểm thử luồng HTTP: upload -> bóc tách -> soát -> lưu thành lần khám.

Bộ bóc tách giả lập được ép trong ``conftest.py`` (ghi đè thẳng vào ``settings``,
vì đặt biến môi trường ở đây là quá muộn — singleton đã dựng xong). Nhờ vậy test
không gọi ra ngoài, không tốn tiền và cho kết quả ổn định.
"""

from __future__ import annotations

import uuid

import pytest
from fastapi.testclient import TestClient

from app.main import app


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as test_client:
        yield test_client


@pytest.fixture(scope="module")
def token(client):
    response = client.post(
        "/api/v1/auth/register",
        json={
            "email": f"ocr-{uuid.uuid4().hex[:8]}@example.com",
            "password": "password123",
            "full_name": "Bác sĩ OCR",
            "role": "doctor",
        },
    )
    assert response.status_code == 201
    return response.json()["access_token"]


@pytest.fixture(scope="module")
def auth(token):
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture(scope="module")
def patient_id(client, auth):
    response = client.post(
        "/api/v1/patients",
        headers=auth,
        json={
            "patient_code": f"OCR{uuid.uuid4().hex[:8]}",
            "full_name": "Bệnh nhân OCR",
            "diagnosis": "Trầm cảm",
        },
    )
    assert response.status_code == 201
    return response.json()["data"]["id"]


def _docx_bytes() -> bytes:
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("PHẦN 3: KHÁM BỆNH")
    document.add_paragraph("Thể trạng: Trung bình")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _png_bytes() -> bytes:
    import fitz

    document = fitz.open()
    page = document.new_page()
    data = page.get_pixmap(dpi=36).tobytes("png")
    document.close()
    return data


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ==================== Bóc tách ====================


def test_upload_mode_returns_draft(client, auth, patient_id):
    response = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=upload",
        headers=auth,
        files={"file": ("phieu.docx", _docx_bytes(), DOCX_MIME)},
    )
    assert response.status_code == 200, response.text
    draft = response.json()["data"]

    assert draft["mode"] == "upload"
    assert draft["provider"] == "stub"
    assert draft["filled_count"] > 0
    assert draft["total_fields"] > 50
    # Dữ liệu phải theo đúng cấu trúc lồng của form.
    assert "general" in draft["data"]
    assert isinstance(draft["data"]["general"], dict)
    # Bệnh nhân trầm cảm -> phải có bảng triệu chứng cơ thể của mẫu F32.
    assert "somatic" in draft["data"]["mental_exam"]


def test_ocr_mode_accepts_image(client, auth, patient_id):
    response = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=ocr",
        headers=auth,
        files={"file": ("scan.png", _png_bytes(), "image/png")},
    )
    assert response.status_code == 200, response.text
    assert response.json()["data"]["mode"] == "ocr"


def test_unsupported_file_rejected(client, auth, patient_id):
    response = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=upload",
        headers=auth,
        files={"file": ("ghi_chu.txt", b"noi dung", "text/plain")},
    )
    assert response.status_code == 422
    assert "Không hỗ trợ" in response.json()["detail"]


def test_empty_file_rejected(client, auth, patient_id):
    response = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=upload",
        headers=auth,
        files={"file": ("rong.docx", b"", DOCX_MIME)},
    )
    assert response.status_code == 422
    assert "rỗng" in response.json()["detail"]


def test_invalid_mode_rejected(client, auth, patient_id):
    response = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=xyz",
        headers=auth,
        files={"file": ("phieu.docx", _docx_bytes(), DOCX_MIME)},
    )
    assert response.status_code == 422


def test_extract_requires_auth(client, patient_id):
    response = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=upload",
        files={"file": ("phieu.docx", _docx_bytes(), DOCX_MIME)},
    )
    assert response.status_code == 401


# ==================== Đọc lại bản nháp & xem file ====================


def test_draft_can_be_reloaded_and_file_served(client, auth, patient_id):
    created = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=upload",
        headers=auth,
        files={"file": ("phieu.docx", _docx_bytes(), DOCX_MIME)},
    ).json()["data"]

    reloaded = client.get(
        f"/api/v1/patients/{patient_id}/exams/extractions/{created['extraction_id']}",
        headers=auth,
    )
    assert reloaded.status_code == 200
    assert reloaded.json()["data"]["data"] == created["data"]

    served = client.get(
        f"/api/v1/patients/{patient_id}/documents/{created['document_id']}/file",
        headers=auth,
    )
    assert served.status_code == 200
    assert served.content[:2] == b"PK"  # docx là file zip


def test_draft_of_another_patient_is_not_reachable(client, auth, patient_id):
    """Bản nháp của bệnh nhân này không được lộ sang hồ sơ bệnh nhân khác."""
    other = client.post(
        "/api/v1/patients",
        headers=auth,
        json={"patient_code": f"OTHER{uuid.uuid4().hex[:6]}", "full_name": "Người khác"},
    ).json()["data"]["id"]

    draft = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=upload",
        headers=auth,
        files={"file": ("phieu.docx", _docx_bytes(), DOCX_MIME)},
    ).json()["data"]

    leaked = client.get(
        f"/api/v1/patients/{other}/exams/extractions/{draft['extraction_id']}",
        headers=auth,
    )
    assert leaked.status_code == 404

    leaked_file = client.get(
        f"/api/v1/patients/{other}/documents/{draft['document_id']}/file", headers=auth
    )
    assert leaked_file.status_code == 404


# ==================== Lưu thành lần khám ====================


def test_save_draft_creates_examination_with_provenance(client, auth, patient_id):
    draft = client.post(
        f"/api/v1/patients/{patient_id}/exams/extract?mode=ocr",
        headers=auth,
        files={"file": ("scan.png", _png_bytes(), "image/png")},
    ).json()["data"]

    data = draft["data"]
    data.setdefault("exam_info", {})["exam_date"] = "2026-05-20"
    data["diagnosis"] = {"definitive": "Bác sĩ đã sửa lại"}

    created = client.post(
        f"/api/v1/patients/{patient_id}/exams",
        headers=auth,
        json={
            "data": data,
            "source": "ocr",
            "document_id": draft["document_id"],
            "extraction_id": draft["extraction_id"],
        },
    )
    assert created.status_code == 201, created.text
    exam = created.json()["data"]

    assert exam["source"] == "ocr"
    assert exam["document_id"] == draft["document_id"]
    # Ngày khám nằm trong data phải được rút ra cột riêng để còn sắp xếp.
    assert exam["exam_date"] == "2026-05-20"
    assert exam["data"]["diagnosis"]["definitive"] == "Bác sĩ đã sửa lại"

    # Lần khám mới phải xuất hiện trong bảng.
    listed = client.get(f"/api/v1/patients/{patient_id}/exams", headers=auth).json()
    assert any(item["id"] == exam["id"] for item in listed["data"])

    # Bản nháp được đánh dấu đã soát, giữ lại bản sửa của bác sĩ.
    reloaded = client.get(
        f"/api/v1/patients/{patient_id}/exams/extractions/{draft['extraction_id']}",
        headers=auth,
    ).json()["data"]
    assert reloaded["data"]["diagnosis"]["definitive"] == "Bác sĩ đã sửa lại"


def test_manual_exam_defaults_to_manual_source(client, auth, patient_id):
    created = client.post(
        f"/api/v1/patients/{patient_id}/exams",
        headers=auth,
        json={"data": {"exam_info": {"exam_date": "2026-06-01"}}},
    )
    assert created.status_code == 201
    assert created.json()["data"]["source"] == "manual"
